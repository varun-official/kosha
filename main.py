from fastapi import FastAPI, File, UploadFile, HTTPException
import numpy as np
import psycopg2
import os
from langchain_core.prompts import ChatPromptTemplate
from langchain_mistralai.chat_models import ChatMistralAI
from langchain_mistralai.embeddings import MistralAIEmbeddings
from PyPDF2 import PdfReader
from docx import Document
import langchain.schema
import httpx
import time
from langchain_community.vectorstores import PGVector
from langchain.chains import RetrievalQA
from pydantic import BaseModel
from dotenv import load_dotenv
import json

from langchain.text_splitter import RecursiveCharacterTextSplitter

# Load environment variables
load_dotenv()

app = FastAPI()

class QueryRequest(BaseModel):
    query: str
    top_k: int = 5

# Initialize Mistral client
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")
DB_CONN_STRING = os.getenv("DATABASE_CONNECTION_STRING")
llm = ChatMistralAI(api_key=MISTRAL_API_KEY)

# Database connection function
def get_db_connection():
    return psycopg2.connect(DB_CONN_STRING)

# Ensure pgvector extension is enabled and update schema
def setup_database():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
    CREATE EXTENSION IF NOT EXISTS vector;
    """)
    conn.commit()
    cur.close()
    conn.close()
setup_database()


def split_text_into_chunks(text, chunk_size=400, chunk_overlap=50):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return text_splitter.split_text(text)

# Initialize MistralAI Embeddings
embedding_model = MistralAIEmbeddings(model="mistral-embed", mistral_api_key=MISTRAL_API_KEY)

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    extracted_text = []

    # Extract text from PDF
    if file.filename.endswith(".pdf"):
        reader = PdfReader(file.file)
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                extracted_text.append((page_num, page_text))
    
    elif file.filename.endswith(".docx"):
        doc = Document(file.file)
        page_num = 1
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text.append((page_num, para.text.strip()))
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload a PDF or DOCX.")

    # Prepare LangChain documents with metadata
    all_chunks = []
    for page_num, page_text in extracted_text:
        cleaned_text = page_text.replace("\n", " ").strip()
        chunks = split_text_into_chunks(cleaned_text)
        for i, chunk in enumerate(chunks):
            all_chunks.append(langchain.schema.Document(
                page_content=chunk,
                metadata={
                    "doc_name": file.filename,
                    "page": page_num,
                    "chunk_index": i
                }
            ))

    # Store using LangChain's PGVector
    PGVector.from_documents(
        documents=all_chunks,
        embedding=embedding_model,
        collection_name="document_embeddings",
        connection_string=DB_CONN_STRING,
    )

    return {"message": "File processed successfully", "filename": file.filename}


def query_from_retrievalqa(qa_chain,query, retries=3, delay=5):
    for attempt in range(retries):
        try:
            return qa_chain({"query": query})
        except httpx.HTTPStatusError as e:
            if "429" in str(e):  # Check if it's a rate limit error
                if attempt < retries - 1:
                    time.sleep(delay)
                else:
                    return "Rate limit exceeded. Please try again later."
            else:
                raise e  # Reraise non-rate limit errors

@app.post("/query")
async def query_embeddings(request: QueryRequest):
    try:
        # PGVector VectorStore
        vectorstore = PGVector(
            connection_string=DB_CONN_STRING,
            collection_name="document_embeddings",
            embedding_function=embedding_model
        )

        # Retriever
        retriever = vectorstore.as_retriever(search_kwargs={"k": request.top_k})

        # Custom prompt
        custom_prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an AI assistant that answers queries strictly based on the given documents.
        - Only use the provided content to generate responses.
        - Do not use external knowledge or speculate.
        - If the provided text does not contain relevant information, respond only with: 'I do not have any information about the query.'
        - Double check the answer before sending

        Query: {question}

        Below is the relevant content extracted from documents:
        {context}
        """)
        ])

        # QA Chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=retriever,
            chain_type="stuff",
            chain_type_kwargs={"prompt": custom_prompt},
            verbose=True,
            return_source_documents=True
        )

        result = query_from_retrievalqa(qa_chain,request.query)

        sources = []
        seen = set()
        for doc in result["source_documents"]:
            meta = doc.metadata
            key = (meta.get("doc_name"), meta.get("page"))
            if key not in seen:
                seen.add(key)
                sources.append({
                    "doc_name": key[0],
                    "page": key[1]
                })
        
        return {
            "query": request.query,
            "response": result["result"],
            "sources":sources
        }
    except Exception as e:
        return {
            "error":str(e)
        }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
